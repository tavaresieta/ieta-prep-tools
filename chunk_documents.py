"""
chunk_documents.py - Divide documentos grandes em pedaços menores

Use este script quando documentos forem muito grandes para processar.

Como usar:
1. Colocar documentos grandes em 'documentos_grandes/'
2. Executar: python chunk_documents.py
3. Chunks aparecerão em 'documentos_chunked/'
4. Mover chunks para 'documents/' para usar nas ferramentas
"""

from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

def chunk_pdf(pdf_path, output_folder, pages_per_chunk=15):
    """Divide PDF em chunks de N páginas"""
    print(f"\n📄 Processando PDF: {pdf_path.name}")
    
    try:
        reader = PdfReader(pdf_path)
        total_pages = len(reader.pages)
        
        print(f"   Total de páginas: {total_pages}")
        
        if total_pages <= pages_per_chunk:
            print(f"   ✅ Documento pequeno - copiando sem dividir")
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            output_file = output_folder / f"{pdf_path.stem}.txt"
            output_file.write_text(text, encoding='utf-8')
            return
        
        num_chunks = (total_pages + pages_per_chunk - 1) // pages_per_chunk
        print(f"   📊 Dividindo em {num_chunks} chunks de ~{pages_per_chunk} páginas")
        
        for chunk_num in range(num_chunks):
            start_page = chunk_num * pages_per_chunk
            end_page = min(start_page + pages_per_chunk, total_pages)
            
            chunk_text = []
            chunk_text.append(f"DOCUMENTO: {pdf_path.stem}")
            chunk_text.append(f"PARTE {chunk_num + 1} de {num_chunks}")
            chunk_text.append(f"Páginas {start_page + 1} a {end_page}")
            chunk_text.append("=" * 70)
            chunk_text.append("")
            
            for page_num in range(start_page, end_page):
                page = reader.pages[page_num]
                chunk_text.append(f"\n--- PÁGINA {page_num + 1} ---\n")
                chunk_text.append(page.extract_text())
            
            output_file = output_folder / f"{pdf_path.stem}_parte{chunk_num + 1:02d}de{num_chunks:02d}.txt"
            output_file.write_text("\n".join(chunk_text), encoding='utf-8')
            
            print(f"      ✅ Parte {chunk_num + 1}/{num_chunks} salva")
        
        print(f"   ✅ PDF completo processado!")
        
    except Exception as e:
        print(f"   ❌ Erro: {str(e)}")


def chunk_word(docx_path, output_folder, paragraphs_per_chunk=100):
    """Divide Word em chunks de N parágrafos"""
    print(f"\n📝 Processando Word: {docx_path.name}")
    
    try:
        doc = DocxDocument(docx_path)
        total_paragraphs = len(doc.paragraphs)
        
        print(f"   Total de parágrafos: {total_paragraphs}")
        
        if total_paragraphs <= paragraphs_per_chunk:
            print(f"   ✅ Documento pequeno - copiando sem dividir")
            text = "\n".join([p.text for p in doc.paragraphs])
            output_file = output_folder / f"{docx_path.stem}.txt"
            output_file.write_text(text, encoding='utf-8')
            return
        
        num_chunks = (total_paragraphs + paragraphs_per_chunk - 1) // paragraphs_per_chunk
        print(f"   📊 Dividindo em {num_chunks} chunks de ~{paragraphs_per_chunk} parágrafos")
        
        for chunk_num in range(num_chunks):
            start_para = chunk_num * paragraphs_per_chunk
            end_para = min(start_para + paragraphs_per_chunk, total_paragraphs)
            
            chunk_text = []
            chunk_text.append(f"DOCUMENTO: {docx_path.stem}")
            chunk_text.append(f"PARTE {chunk_num + 1} de {num_chunks}")
            chunk_text.append(f"Parágrafos {start_para + 1} a {end_para}")
            chunk_text.append("=" * 70)
            chunk_text.append("")
            
            for para_num in range(start_para, end_para):
                para_text = doc.paragraphs[para_num].text.strip()
                if para_text:
                    chunk_text.append(para_text)
                    chunk_text.append("")
            
            output_file = output_folder / f"{docx_path.stem}_parte{chunk_num + 1:02d}de{num_chunks:02d}.txt"
            output_file.write_text("\n".join(chunk_text), encoding='utf-8')
            
            print(f"      ✅ Parte {chunk_num + 1}/{num_chunks} salva")
        
        print(f"   ✅ Word completo processado!")
        
    except Exception as e:
        print(f"   ❌ Erro: {str(e)}")


def main():
    print("=" * 70)
    print("DIVISOR DE DOCUMENTOS GRANDES - IETA")
    print("=" * 70)
    
    input_folder = Path("documentos_grandes")
    output_folder = Path("documentos_chunked")
    
    input_folder.mkdir(exist_ok=True)
    output_folder.mkdir(exist_ok=True)
    
    pdf_files = list(input_folder.glob("*.pdf"))
    word_files = list(input_folder.glob("*.docx"))
    
    total_files = len(pdf_files) + len(word_files)
    
    if total_files == 0:
        print("\n⚠️  Nenhum documento encontrado!")
        print(f"   Coloque PDFs ou Words na pasta: {input_folder.absolute()}")
        print("\n💡 Dica: Crie a pasta 'documentos_grandes' e adicione arquivos lá.")
        return
    
    print(f"\n📁 Encontrados {total_files} documentos:")
    print(f"   - PDFs: {len(pdf_files)}")
    print(f"   - Words: {len(word_files)}")
    
    if pdf_files:
        print("\n" + "=" * 70)
        print("PROCESSANDO PDFs")
        print("=" * 70)
        
        for pdf_file in pdf_files:
            chunk_pdf(pdf_file, output_folder, pages_per_chunk=15)
    
    if word_files:
        print("\n" + "=" * 70)
        print("PROCESSANDO WORDS")
        print("=" * 70)
        
        for word_file in word_files:
            chunk_word(word_file, output_folder, paragraphs_per_chunk=100)
    
    output_files = list(output_folder.glob("*.txt"))
    
    print("\n" + "=" * 70)
    print("PROCESSAMENTO CONCLUÍDO!")
    print("=" * 70)
    print(f"\n✅ {len(output_files)} arquivos criados em: {output_folder.absolute()}")
    print("\n📋 Próximos passos:")
    print("   1. Revisar arquivos em 'documentos_chunked/'")
    print("   2. Mover para 'documents/' para usar nas ferramentas")
    print("   3. Clicar em 'Recarregar Documentos' no app")


if __name__ == "__main__":
    main()